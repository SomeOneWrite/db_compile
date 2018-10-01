import os
import sqlite3
from docx import Document
import json


query_create_tables = [
    'CREATE TABLE dirs ( id integer primary key autoincrement, parent_id integer references dirs(id) on delete cascade on update cascade, name text )',
    'CREATE TABLE collections ( id integer primary key autoincrement, dir_id integer references dirs(id) on delete cascade on update cascade, type integer not null, name text, techpart text )',
    'CREATE TABLE captions ( id integer primary key autoincrement, collection_id integer not null references collections(id) on delete cascade on update cascade, parent_id integer references captions(id) on delete cascade on update cascade, name text )',
    'CREATE TABLE unit_positions ( id text not null, name text, unit text, cost_workers real, cost_machines real, cost_drivers real, cost_materials real, mass real, caption_id integer not null references captions(id) on delete cascade on update cascade )',
    'CREATE TABLE options ( name text not null, val text, primary key(name) )',
    'CREATE TABLE materials ( id text not null, name text, unit text, price real, price_vacantion real, caption_id integer references captions(id) on delete set null on update cascade, primary key(id) )',
    'CREATE TABLE machines ( id text not null, name text, unit text, price real, price_driver real, caption_id integer references captions(id) on delete set null on update cascade, primary key(id) )',
    'CREATE TABLE transports ( id text not null, name text, unit text, price real, type integer check(type >= 1 and type <= 4), caption_id integer references captions(id) on delete set null on update cascade, primary key(id) )',
    'CREATE TABLE workers ( id text not null, name text, unit text, price real, caption_id integer references captions(id) on delete set null on update cascade, primary key(id) )',
]


class FileModel:
    def __init__(self, db_name):
        self.db_name = db_name
        self.tables = {}
        self.init_db()
        self.last_caption_id = 1
        self.last_collection_id = 1
        self.last_dir_id = 1
        self.last_unit_position_id = 1
        self.count = 0

    def init_db(self):
        if os.path.exists(self.db_name):
            os.remove(self.db_name)


    def insert_caption(self, collection_id: int, parent_id: int, name: str):
        # print('insert caption: coll_id = {}     parent_id = {}     name = {}'.format(collection_id, parent_id, name))
        self.tables[self.last_caption_id] = [collection_id, parent_id, name]
        self.last_caption_id += 1
        return self.last_caption_id

    def insert_collection(self, dir_id: int, type: int, name: str, tech_part: str):
        self.tables[self.last_collection_id] = [dir_id, type, name, tech_part]
        self.last_collection_id += 1
        return self.last_collection_id

    def insert_dir(self, parent_id: int, name: str):
        self.tables[self.last_dir_id] = [parent_id, name]
        self.last_dir_id += 1
        return self.last_dir_id

    def insert_unit_position(self, id: str, name:str, unit:str, cost_workers:str, cost_machines: str, cost_drivers: str,
                             cost_materials: str, caption_id: int):
        self.tables[self.last_unit_position_id] = [id, name, unit, cost_workers, cost_machines, cost_drivers, cost_materials, caption_id]
        self.last_unit_position_id += 1
        print("id {} : {}".format(id, name))
        return self.last_unit_position_id

    def commit(self):
        with open(self.db_name, 'w', encoding='utf-8') as f:
            json.dump(self.tables, f, ensure_ascii=False)

    def __del__(self):
        self.commit()


class SqlModel:
    def __init__(self, db_name):
        self.db_name = db_name
        self.init_db()

    def init_db(self):
        if os.path.exists(self.db_name):
            os.remove(self.db_name)
        self.db_connection = sqlite3.connect(self.db_name)
        self.db_cursor = self.db_connection.cursor()
        for query_create in query_create_tables:
            try:
                self.db_cursor.execute(query_create)
            except sqlite3.DatabaseError as err:
                print("Error: ", err)
            else:
                self.db_connection.commit()
        print("Database and tables created")

    def insert_caption(self, collection_id: int, parent_id: int, name: str):
        # print('insert caption: coll_id = {}     parent_id = {}     name = {}'.format(collection_id, parent_id, name))
        return self.db_cursor.execute('insert into captions (collection_id, parent_id, name) values(?, ?, ?)',
                                      (collection_id, parent_id, name)).lastrowid

    def insert_collection(self, dir_id: int, type: int, name: str, tech_part: str):
    #     print('insert collection: dir_id = {} name = {}'.format(dir_id, name))
        return self.db_cursor.execute('insert into collections (dir_id, type, name, techpart) values(?, ?, ?, ?)',
                                      (dir_id, type, name, tech_part)).lastrowid

    def insert_dir(self, parent_id: int, name: str):
        # print('insert parent_id: parent_id = {}      name = {}'.format(parent_id, name))
        return self.db_cursor.execute('insert into dirs(parent_id, name) values(?, ?)', (parent_id, name)).lastrowid

    def insert_unit_position(self, id: str, name:str, unit:str, cost_workers:str, cost_machines: str, cost_drivers: str,
                             cost_materials: str, caption_id: int):
        print('insert unit_position: id = {} w = {}  m = {} d = {} m = {}'.format(id, cost_workers, cost_machines, cost_drivers, cost_materials))
        return \
            self.db_cursor.execute("insert into unit_positions"
                                      "(id, name, unit, cost_workers, cost_machines, cost_drivers, cost_materials, caption_id) "
                                      "values(?, ?, ?, ?, ?, ?, ?, ?)",
                                   (id, name, unit, cost_workers, cost_machines, cost_drivers, cost_materials, caption_id)).lastrowid
    def commit(self):
        self.db_connection.commit()

    def __del__(self):
        self.commit()

